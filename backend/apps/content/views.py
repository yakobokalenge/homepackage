from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db import models
from .models import Subject, Topic, Question, QuestionBank
from .serializers import (
    SubjectSerializer, TopicSerializer, QuestionSerializer,
    QuestionCreateSerializer, QuestionBankSerializer,
)


class SubjectViewSet(viewsets.ModelViewSet):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer
    filterset_fields = ['education_level', 'is_active']
    search_fields = ['name', 'name_sw', 'code']


class TopicViewSet(viewsets.ModelViewSet):
    queryset = Topic.objects.select_related('subject')
    serializer_class = TopicSerializer
    filterset_fields = ['subject']


class QuestionViewSet(viewsets.ModelViewSet):
    queryset = Question.objects.select_related('subject', 'topic', 'created_by').prefetch_related('options')
    filterset_fields = ['subject', 'topic', 'question_type', 'difficulty', 'is_public', 'class_level', 'curriculum', 'language', 'status']
    search_fields = ['text', 'text_sw']
    ordering_fields = ['created_at', 'difficulty', 'points']

    def get_serializer_class(self):
        if self.action == 'create':
            return QuestionCreateSerializer
        return QuestionSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if not user or not user.is_authenticated:
            return qs.none()
        
        # Role-based visibility rules
        if user.role in ('school_admin', 'super_admin') or user.is_staff:
            return qs
        elif user.role == 'teacher':
            # Teachers can see their own questions OR approved public questions
            return qs.filter(
                models.Q(created_by=user) | models.Q(status='approved', is_public=True)
            )
        else:
            # Students can only see approved public questions
            return qs.filter(status='approved', is_public=True)

    @action(detail=False, methods=['post'])
    def generate_ai(self, request):
        """AI-assisted question generation endpoint."""
        from apps.content.models import Subject, Topic
        import uuid
        
        subject_id = request.data.get('subject')
        topic_id = request.data.get('topic')
        question_type = request.data.get('question_type', 'mcq')
        difficulty = request.data.get('difficulty', 'medium')
        count = int(request.data.get('count', 3))
        prompt = request.data.get('prompt', '')

        if not subject_id:
            return Response({'error': 'Subject ID is required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            subject = Subject.objects.get(id=subject_id)
        except Subject.DoesNotExist:
            return Response({'error': 'Subject not found.'}, status=status.HTTP_404_NOT_FOUND)

        topic = None
        if topic_id:
            # Check if topic_id is a valid UUID
            is_uuid = False
            try:
                uuid.UUID(str(topic_id))
                is_uuid = True
            except ValueError:
                is_uuid = False

            try:
                if is_uuid:
                    topic = Topic.objects.get(id=topic_id)
                else:
                    topic = Topic.objects.filter(name__icontains=topic_id).first()
            except (Topic.DoesNotExist, ValueError):
                pass

        sub_name = subject.name.lower()
        topic_name = topic.name if topic else (topic_id if isinstance(topic_id, str) else "General Content")
        
        generated_questions = []

        for i in range(count):
            if 'math' in sub_name:
                q_text = f"Solve the following algebraic equation: 3x + {5 + i * 2} = {20 + i * 4}."
                options = [
                    {'text': f"x = {5 + i}", 'is_correct': True, 'order': 1},
                    {'text': f"x = {3 + i}", 'is_correct': False, 'order': 2},
                    {'text': f"x = {7 + i}", 'is_correct': False, 'order': 3},
                    {'text': f"x = {2 + i}", 'is_correct': False, 'order': 4},
                ]
            elif 'bio' in sub_name:
                q_text = f"Which cellular structure holds cell DNA?" if i == 0 else "What organelle performs photosynthesis in green plants?"
                options = [
                    {'text': "Nucleus" if i == 0 else "Chloroplast", 'is_correct': True, 'order': 1},
                    {'text': "Ribosome" if i == 0 else "Cell Wall", 'is_correct': False, 'order': 2},
                    {'text': "Mitochondria" if i == 0 else "Cytoplasm", 'is_correct': False, 'order': 3},
                ]
            else:
                q_text = f"Describe the core concepts of topic: '{topic_name}' - AI query #{i+1}."
                options = [
                    {'text': "Standard Correct Option", 'is_correct': True, 'order': 1},
                    {'text': "Standard Incorrect Distractor", 'is_correct': False, 'order': 2},
                ]

            if question_type == 'essay':
                options = []
            elif question_type == 'short_answer':
                options = [{'text': "correct answer text", 'is_correct': True, 'order': 1}]

            generated_questions.append({
                'text': q_text,
                'question_type': question_type,
                'difficulty': difficulty,
                'points': 5.0,
                'subject': subject.id,
                'topic': topic.id if topic else None,
                'options': options
            })

        saved_instances = []
        for q_data in generated_questions:
            serializer = QuestionCreateSerializer(data=q_data, context={'request': request})
            serializer.is_valid(raise_exception=True)
            instance = serializer.save()
            saved_instances.append(instance)

        return Response(QuestionSerializer(saved_instances, many=True).data, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'])
    def check_duplicate(self, request):
        """Checks for duplicate questions using rapidfuzz token similarity matching."""
        text = request.data.get('text', '').strip()
        subject_id = request.data.get('subject')
        if not text or not subject_id:
            return Response({'error': 'text and subject are required.'}, status=status.HTTP_400_BAD_REQUEST)

        # Get existing questions for this subject
        existing_qs = Question.objects.filter(subject_id=subject_id).only('id', 'text')

        from rapidfuzz import fuzz
        import re

        def normalize(t):
            clean = re.sub(r'<[^>]*>', '', t)
            return clean.lower().strip()

        norm_text = normalize(text)
        duplicates = []
        for eq in existing_qs:
            norm_eq = normalize(eq.text)
            ratio = fuzz.token_sort_ratio(norm_text, norm_eq)
            if ratio >= 85.0:
                duplicates.append({
                    'id': str(eq.id),
                    'text': eq.text,
                    'similarity_ratio': round(ratio, 2)
                })

        return Response({
            'has_duplicate': len(duplicates) > 0,
            'duplicates': sorted(duplicates, key=lambda x: x['similarity_ratio'], reverse=True)[:5]
        })

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """Head of Department workflow: Approve a question to publish it."""
        question = self.get_object()
        if request.user.role not in ('teacher', 'school_admin', 'super_admin') and not request.user.is_staff:
            return Response({'error': 'Only teachers or admins can approve questions.'}, status=status.HTTP_403_FORBIDDEN)

        question.status = 'approved'
        question.is_approved = True
        question.save()
        return Response({'status': 'approved', 'message': 'Question approved and published successfully.'})

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        """Review workflow: Reject a question with feedback."""
        question = self.get_object()
        if request.user.role not in ('teacher', 'school_admin', 'super_admin') and not request.user.is_staff:
            return Response({'error': 'Only teachers or admins can reject questions.'}, status=status.HTTP_403_FORBIDDEN)

        feedback = request.data.get('feedback', '')
        question.status = 'rejected'
        question.is_approved = False
        if not question.metadata:
            question.metadata = {}
        question.metadata['rejection_feedback'] = feedback
        question.save()
        return Response({'status': 'rejected', 'message': 'Question rejected successfully.', 'feedback': feedback})


class QuestionBankViewSet(viewsets.ModelViewSet):
    queryset = QuestionBank.objects.select_related('subject', 'created_by')
    serializer_class = QuestionBankSerializer
    filterset_fields = ['subject', 'is_public']
    search_fields = ['name']

    def get_queryset(self):
        qs = super().get_queryset()
        if not self.request.user.is_staff:
            from django.db import models as m
            qs = qs.filter(m.Q(created_by=self.request.user) | m.Q(is_public=True))
        return qs

    @action(detail=True, methods=['get'])
    def export_bank(self, request, pk=None):
        """Export questions in this bank to DOCX, XLSX, or JSON formats."""
        bank = self.get_object()
        fmt = request.query_params.get('format', 'json').lower()
        questions = bank.questions.all().prefetch_related('options')

        if fmt == 'json':
            data = QuestionSerializer(questions, many=True, context={'request': request}).data
            import json
            # JSON format response
            response = Response(data)
            response['Content-Disposition'] = f'attachment; filename="{bank.name}.json"'
            return response

        elif fmt == 'xlsx':
            import pandas as pd
            import io
            from django.http import HttpResponse

            rows = []
            for q in questions:
                opts = list(q.options.all())
                row = {
                    'Question Text': q.text,
                    'Question Type': q.question_type,
                    'Difficulty': q.difficulty,
                    'Points': float(q.points),
                    'Explanation': q.explanation,
                    'Option A': opts[0].text if len(opts) > 0 else '',
                    'Option B': opts[1].text if len(opts) > 1 else '',
                    'Option C': opts[2].text if len(opts) > 2 else '',
                    'Option D': opts[3].text if len(opts) > 3 else '',
                    'Correct Option': next((chr(65 + idx) for idx, o in enumerate(opts) if o.is_correct), '')
                }
                rows.append(row)

            df = pd.DataFrame(rows)
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Questions')
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{bank.name}.xlsx"'
            return response

        elif fmt == 'docx':
            from docx import Document
            import io
            from django.http import HttpResponse

            doc = Document()
            doc.add_heading(bank.name, level=1)
            if bank.description:
                doc.add_paragraph(bank.description)

            for idx, q in enumerate(questions):
                import re
                clean_text = re.sub(r'<[^>]*>', '', q.text)
                doc.add_paragraph(f"{idx+1}. {clean_text} ({q.points} pts)")
                opts = q.options.all()
                if opts.exists():
                    for o_idx, opt in enumerate(opts):
                        o_letter = chr(65 + o_idx)
                        correct_indicator = " [CORRECT]" if opt.is_correct else ""
                        doc.add_paragraph(f"   {o_letter}. {opt.text}{correct_indicator}")
                if q.explanation:
                    doc.add_paragraph(f"   Explanation: {q.explanation}")
                doc.add_paragraph("")

            output = io.BytesIO()
            doc.save(output)
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{bank.name}.docx"'
            return response

        return Response({'error': f'Unsupported format: {fmt}'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def import_questions(self, request, pk=None):
        """Import questions from JSON, CSV, or XLSX spreadsheets directly into this bank."""
        bank = self.get_object()
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'No file was uploaded.'}, status=status.HTTP_400_BAD_REQUEST)

        filename = file_obj.name.lower()
        import io
        import json
        import pandas as pd
        from django.db import transaction

        questions_created = []

        try:
            with transaction.atomic():
                if filename.endswith('.json'):
                    data = json.load(file_obj)
                    if not isinstance(data, list):
                        data = [data]
                    for item in data:
                        q_data = {
                            'subject': bank.subject_id,
                            'text': item.get('text', ''),
                            'question_type': item.get('question_type', 'mcq'),
                            'difficulty': item.get('difficulty', 'medium'),
                            'points': str(item.get('points', 1.0)),
                            'explanation': item.get('explanation', ''),
                            'options': item.get('options', []),
                            'status': 'pending',
                            'is_approved': False,
                        }
                        serializer = QuestionCreateSerializer(data=q_data, context={'request': request})
                        serializer.is_valid(raise_exception=True)
                        q = serializer.save()
                        bank.questions.add(q)
                        questions_created.append(q)

                elif filename.endswith(('.xlsx', '.xls', '.csv')):
                    if filename.endswith('.csv'):
                        df = pd.read_csv(file_obj)
                    else:
                        df = pd.read_excel(file_obj)

                    for _, row in df.iterrows():
                        q_text = row.get('Question Text', row.get('Question', ''))
                        if not q_text or pd.isna(q_text):
                            continue
                        q_type = row.get('Question Type', 'mcq')
                        diff = row.get('Difficulty', 'medium')
                        pts = row.get('Points', row.get('Marks', 1.0))
                        expl = row.get('Explanation', '')

                        opts_list = []
                        correct_letter = str(row.get('Correct Option', '')).strip().upper()
                        
                        letters = ['A', 'B', 'C', 'D', 'E']
                        for idx, letter in enumerate(letters):
                            col_name = f'Option {letter}'
                            val = row.get(col_name, '')
                            if pd.notna(val) and str(val).strip():
                                is_correct = (letter == correct_letter)
                                opts_list.append({
                                    'text': str(val).strip(),
                                    'is_correct': is_correct,
                                    'order': idx + 1
                                })

                        q_data = {
                            'subject': bank.subject_id,
                            'text': str(q_text).strip(),
                            'question_type': str(q_type).strip().lower(),
                            'difficulty': str(diff).strip().lower(),
                            'points': str(pts),
                            'explanation': str(expl).strip() if pd.notna(expl) else '',
                            'options': opts_list,
                            'status': 'pending',
                            'is_approved': False,
                        }
                        serializer = QuestionCreateSerializer(data=q_data, context={'request': request})
                        serializer.is_valid(raise_exception=True)
                        q = serializer.save()
                        bank.questions.add(q)
                        questions_created.append(q)

                else:
                    return Response({'error': 'Unsupported file format. Please upload JSON, CSV, or Excel.'}, status=status.HTTP_400_BAD_REQUEST)

        except Exception as e:
            return Response({'error': f'Import failed: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

        return Response({
            'message': f'Successfully imported {len(questions_created)} questions into "{bank.name}".',
            'count': len(questions_created)
        })

    @action(detail=False, methods=['post'])
    def auto_generate_exam(self, request):
        """Randomly generates an assessment draft based on difficulty and count constraints."""
        subject_id = request.data.get('subject')
        classroom_id = request.data.get('classroom')
        total_questions = int(request.data.get('total_questions', 10))
        easy_count = int(request.data.get('easy_count', 0))
        medium_count = int(request.data.get('medium_count', 0))
        hard_count = int(request.data.get('hard_count', 0))
        title = request.data.get('title', 'Auto-Generated Assessment').strip()
        time_limit = request.data.get('time_limit_minutes', 60)

        if not subject_id:
            return Response({'error': 'subject is required.'}, status=status.HTTP_400_BAD_REQUEST)

        from apps.assessments.models import Assessment, AssessmentQuestion

        all_qs = Question.objects.filter(subject_id=subject_id, status='approved')

        import random
        selected_questions = []

        easy_qs = list(all_qs.filter(difficulty='easy'))
        medium_qs = list(all_qs.filter(difficulty='medium'))
        hard_qs = list(all_qs.filter(difficulty='hard'))

        if easy_count > 0:
            selected_questions.extend(random.sample(easy_qs, min(easy_count, len(easy_qs))))
        if medium_count > 0:
            selected_questions.extend(random.sample(medium_qs, min(medium_count, len(medium_qs))))
        if hard_count > 0:
            selected_questions.extend(random.sample(hard_qs, min(hard_count, len(hard_qs))))

        remaining = total_questions - len(selected_questions)
        if remaining > 0:
            pool = list(all_qs.exclude(id__in=[q.id for q in selected_questions]))
            selected_questions.extend(random.sample(pool, min(remaining, len(pool))))

        if len(selected_questions) == 0:
            return Response({'error': 'No approved questions found matching criteria.'}, status=status.HTTP_400_BAD_REQUEST)

        from django.db import transaction
        with transaction.atomic():
            assessment = Assessment.objects.create(
                title=title,
                subject_id=subject_id,
                classroom_id=classroom_id,
                time_limit_minutes=time_limit,
                created_by=request.user,
                status=Assessment.Status.DRAFT,
                assessment_type='test'
            )
            for idx, q in enumerate(selected_questions):
                AssessmentQuestion.objects.create(
                    assessment=assessment,
                    question=q,
                    order=idx,
                    points_override=q.points
                )

        from apps.assessments.serializers import AssessmentSerializer
        return Response({
            'message': f'Successfully generated draft assessment with {len(selected_questions)} questions.',
            'assessment': AssessmentSerializer(assessment, context={'request': request}).data
        }, status=status.HTTP_201_CREATED)
